from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Colours ──────────────────────────────────────────────────────────────
GREEN  = RGBColor(0x2D, 0x7A, 0x3A)
GOLD   = RGBColor(0xD4, 0xAF, 0x37)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xC0, 0x39, 0x2B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_border(cell, sides=('top','bottom','left','right'), sz=6, color='CCCCCC'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def h(level, text):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = GREEN
        p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'2D7A3A')
        pBdr.append(bot); pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = DARK
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11); run.font.bold = True; run.font.color.rgb = GREY
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)

def body(text, bold=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)

def info_box(title, lines, bg='E8F5E9', border='2D7A3A'):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    set_cell_bg(cell, bg); add_border(cell, color=border, sz=8)
    p = cell.paragraphs[0]
    if title:
        r = p.add_run(title + '\n'); r.font.bold = True; r.font.size = Pt(10.5)
    for line in lines:
        r = p.add_run('• ' + line + '\n'); r.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'),'DDDDDD')
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def code_block(lines):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    set_cell_bg(cell, '1A1A1A'); add_border(cell, color='2D7A3A', sz=8)
    p = cell.paragraphs[0]
    for line in lines:
        r = p.add_run(line + '\n'); r.font.name = 'Courier New'; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xA8,0xFF,0xC4) if line.startswith('#') else WHITE
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()

# ════════════ COVER ════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('IGO PROTEIN CUTS')
r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = GREEN
p.paragraph_format.space_before = Pt(36); p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('WEBSITE AUDIT, COMPETITIVE ANALYSIS & IMPROVEMENT ROADMAP')
r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = DARK
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared: May 30, 2026  |  Website: igo-protien-cut.vercel.app  |  Compared vs: tendercuts.in')
r.font.size = Pt(9); r.font.color.rgb = GREY
p.paragraph_format.space_after = Pt(20)

divider()

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val, bg, tc) in enumerate([
    ('BUGS FOUND', '12 Issues', 'FFE0E0', 'C0392B'),
    ('FEATURE GAPS vs TenderCuts', '8 Gaps', 'FFF3CD', 'D4AF37'),
    ('DEPLOYMENT', 'Hostinger Ready', 'E8F5E9', '2D7A3A'),
]):
    cell = tbl.cell(0,i)
    set_cell_bg(cell, bg); add_border(cell, color=tc)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(val + '\n'); r.font.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(int(tc[0:2],16), int(tc[2:4],16), int(tc[4:6],16))
    r2 = p.add_run(label); r2.font.size = Pt(8); r2.font.bold = True; r2.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()
doc.add_page_break()

# ════════════ SECTION 1 ════════════
h(1, '1. Website Audit Overview')
body('URL Audited: https://igo-protien-cut.vercel.app/  [NOTE: "protien" is a typo — should be "protein"]', bold=True, color=RED)
body('Tech Stack: React 19 + Vite + TypeScript + TailwindCSS v4 + Framer Motion + Supabase + React Router v7')
body('Current Hosting: Vercel  |  Target: Hostinger with custom domain')
body('')
body('The website has a strong creative foundation — modern glassmorphism design, cohesive green/gold brand colours, rich animations, and high-quality content sections. However, several critical functional gaps and bugs prevent it from being a fully operational e-commerce platform.')

h(2, '1.1 Bugs & Issues Found')

bugs = [
    ('B1','CRITICAL','URL Typo in Vercel project name','"protien" instead of "protein" — visible in URL bar, hurts SEO and brand trust'),
    ('B2','CRITICAL','Blank white area at page bottom','Lenis smooth scroll adds ghost height — page appears blank after Newsletter section'),
    ('B3','HIGH','BrandNarrative section not rendered','Component is imported in App.tsx but never added to the Home JSX — entire section is hidden'),
    ('B4','HIGH','Products stored in localStorage only','All product data lives in browser localStorage — lost when cache is cleared, not a real DB'),
    ('B5','HIGH','No checkout / payment gateway','Cart works but no payment flow — users cannot actually place an order or pay'),
    ('B6','HIGH','No customer login or accounts','Navbar login icon leads nowhere — no sign-up, no order history, no saved addresses'),
    ('B7','MEDIUM','App Store / Play Store buttons are dead links','Both buttons have no href — misleads users into thinking an app exists'),
    ('B8','MEDIUM','No SEO meta tags','Missing og:title, og:image, og:description, robots.txt, sitemap.xml — Google cannot index properly'),
    ('B9','MEDIUM','Fake phone number in footer','+91 9421-IGO-000 is a placeholder — breaks trust and usability'),
    ('B10','MEDIUM','Admin panel has no access control','Admin routes /admin/* require no authentication — anyone who knows the URL can access it'),
    ('B11','LOW','Twitter and WhatsApp footer links are "#"','Dead links — should link to real profiles or be removed'),
    ('B12','LOW','No custom 404 error page','Unknown URLs silently redirect to home instead of showing a helpful 404 page'),
]

tbl = doc.add_table(rows=1+len(bugs), cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, ht in enumerate(['ID','Severity','Bug','Description']):
    cell = tbl.cell(0,i)
    set_cell_bg(cell, '2D7A3A'); add_border(cell, color='2D7A3A')
    p = cell.paragraphs[0]; r = p.add_run(ht)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)

sev_c = {'CRITICAL':'FFD7D7','HIGH':'FFE8CC','MEDIUM':'FFF9CC','LOW':'E8F5E9'}
sev_t = {'CRITICAL': RED, 'HIGH': RGBColor(0xC0,0x6A,0x00), 'MEDIUM': RGBColor(0xA0,0x7A,0x00), 'LOW': GREEN}
for ri, (bid, sev, name, desc) in enumerate(bugs, 1):
    for ci, text in enumerate([bid, sev, name, desc]):
        cell = tbl.cell(ri, ci)
        set_cell_bg(cell, 'FFFFFF' if ri%2==0 else 'F9F9F9'); add_border(cell, color='DDDDDD', sz=4)
        p = cell.paragraphs[0]; r = p.add_run(text)
        r.font.size = Pt(9.5); r.font.bold = (ci < 3)
        if ci == 1: r.font.color.rgb = sev_t.get(sev, DARK)
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
doc.add_page_break()

# ════════════ SECTION 2 ════════════
h(1, '2. Competitive Analysis: IGO Protein Cuts vs TenderCuts')
body('TenderCuts (tendercuts.in) is the benchmark for fresh meat e-commerce in Tamil Nadu. Below is a feature-by-feature comparison:')
doc.add_paragraph()

compare = [
    ('Feature','IGO Protein Cuts','TenderCuts','Priority'),
    ('Design / UI Quality','Premium glassmorphism, modern','Clean, conversion-focused','IGO leads'),
    ('Product Catalog','Static (localStorage)','Live DB with real images','Fix — HIGH'),
    ('Checkout & Payment','No payment gateway','Razorpay fully integrated','CRITICAL'),
    ('User Login / Accounts','No customer accounts','Full login + order history','HIGH'),
    ('Order Tracking','QR mockup only','Real-time SMS + app tracking','HIGH'),
    ('Mobile Responsiveness','Layout breaks on small screens','Fully responsive + PWA','HIGH'),
    ('iOS / Android App','Placeholder buttons only','Live on App Store & Play Store','MEDIUM'),
    ('Prime / Elite Membership','IGO Prime section (no payment)','TenderCuts Elite (live billing)','MEDIUM'),
    ('SEO Optimisation','No meta tags, no sitemap','Full SEO + schema markup','HIGH'),
    ('Search','Local filter only','Full-text + voice search','MEDIUM'),
    ('Reviews & Ratings','Static "12,000+" claim','Real verified post-order reviews','HIGH'),
    ('Blog / Content','Static section (3 articles)','CMS-driven live blog','LOW'),
    ('B2B Portal','Visual section only','Working wholesale login','MEDIUM'),
    ('Delivery Pincode Check','Works (hardcoded list)','Live API-based check','MEDIUM'),
    ('Admin Panel','Functional admin exists','Full ops dashboard','IGO advantage'),
    ('Traceability / QR','Unique QR concept (strong USP)','No equivalent feature','IGO leads'),
]

tbl2 = doc.add_table(rows=len(compare), cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
for ri, row in enumerate(compare):
    is_hdr = (ri == 0)
    for ci, text in enumerate(row):
        cell = tbl2.cell(ri, ci)
        if is_hdr: set_cell_bg(cell, '1A1A1A')
        elif ri%2==0: set_cell_bg(cell, 'F5F5F5')
        else: set_cell_bg(cell, 'FFFFFF')
        add_border(cell, color='CCCCCC', sz=4)
        p = cell.paragraphs[0]; r = p.add_run(text)
        r.font.size = Pt(9.5); r.font.bold = is_hdr
        if is_hdr: r.font.color.rgb = WHITE
        elif ci == 3:
            if 'CRITICAL' in text: r.font.color.rgb = RED
            elif 'HIGH' in text: r.font.color.rgb = RGBColor(0xC0,0x6A,0x00)
            elif 'leads' in text or 'advantage' in text: r.font.color.rgb = GREEN
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()
doc.add_page_break()

# ════════════ SECTION 3 ════════════
h(1, '3. Improvement Plan — Phase by Phase')

h(2, 'Phase 1: Immediate Bug Fixes (Week 1)')
info_box('Fix These First — They Break User Trust & Experience', [
    'Fix blank footer: add Lenis CSS fix so scroll stops at actual page end',
    'Add <BrandNarrative /> to Home component in App.tsx (line after <Blog />)',
    'Replace +91 9421-IGO-000 with real phone number in Footer.tsx',
    'Fix or remove dead Twitter (#) and WhatsApp (#) links in footer',
    'Remove or link App Store / Play Store buttons to real app pages',
    'Add a custom 404.tsx page and route it in App.tsx',
    'Protect all /admin/* routes with Supabase session + role check',
], bg='FFE8E8', border='C0392B')

h(2, 'Phase 2: Core E-commerce (Weeks 2–4)')
body('2A — Payment Integration (Razorpay)', bold=True)
bullet('Create src/pages/Checkout.tsx with: address form, slot picker, cart summary, payment button')
bullet('Integrate Razorpay JS SDK — razorpay.open() on button click, verify signature on backend')
bullet('Use existing Express API in /api folder to handle Razorpay order creation + verification')
bullet('On success: save order to Supabase orders table, send confirmation email via Resend')

body('2B — Customer Authentication (Supabase Auth)', bold=True)
bullet('Enable Email + Google OAuth in Supabase dashboard')
bullet('Connect existing AuthModal.tsx to supabase.auth.signInWithOtp() / signInWithOAuth()')
bullet('Create src/pages/Account.tsx — order history, saved addresses, wishlist')
bullet('Protect checkout page — redirect to login if not authenticated')

body('2C — Real Product Database (Supabase)', bold=True)
bullet('Create products table in Supabase (schema in supabase_setup.sql already exists)')
bullet('Replace localStorage productStore.ts with Supabase client queries')
bullet('Connect admin Product Management to write/update to Supabase products table')
bullet('Add Supabase Storage bucket for product images with CDN URLs')

body('2D — Order Status & Tracking', bold=True)
bullet('Create orders table with status enum: placed, processing, out_for_delivery, delivered')
bullet('Connect the Traceability QR code feature to real order data')
bullet('Add MSG91 or Twilio for SMS delivery updates to customers')

h(2, 'Phase 3: UX, Performance & SEO (Weeks 3–5)')
body('3A — Mobile Fixes', bold=True)
bullet('Test every section at 375px (iPhone SE), 390px (iPhone 15), 768px (iPad)')
bullet('Fix CustomBoxBuilder: Smart Cooler sidebar should become a bottom drawer on mobile')
bullet('Fix CategoryGrid icons overflow on small screens')
bullet('Ensure all text is readable (min 16px on mobile to prevent iOS zoom)')

body('3B — Performance', bold=True)
bullet('Add loading="lazy" to all <img> tags in product cards and sections')
bullet('Convert all JPEG/PNG section images to WebP format (saves 25-35% file size)')
bullet('Add React.lazy / Suspense to heavy sections (BrandNarrative, Traceability, 3D components)')
bullet('Target Lighthouse scores: Performance > 85, SEO > 95, Accessibility > 90')

body('3C — SEO Implementation', bold=True)
bullet('Install react-helmet-async: npm install react-helmet-async')
bullet('Add <HelmetProvider> in main.tsx, add per-page <Helmet> with title/description/og tags')
bullet('Create public/robots.txt (allow all crawlers, link to sitemap)')
bullet('Generate public/sitemap.xml listing: /, /blog, /products, /traceability')
bullet('Add JSON-LD schema for LocalBusiness, Product, and Review types')
bullet('Register and verify domain on Google Search Console')

h(2, 'Phase 4: Advanced Features (Month 2)')
bullet('Launch IGO Prime with real Razorpay subscription billing (monthly/annual plans)')
bullet('Build real B2B wholesale portal with dedicated login, bulk pricing tiers, invoice PDF download')
bullet('Add post-delivery ratings & reviews system (email trigger after order delivered)')
bullet('Add Algolia or Supabase full-text search with autocomplete')
bullet('Build recipe detail pages linked to specific product SKUs (drives upsell)')
bullet('Add multi-language toggle (Tamil / English) for regional appeal')
bullet('PWA support: service worker + offline cache for return visitors')

doc.add_paragraph()
doc.add_page_break()

# ════════════ SECTION 4 ════════════
h(1, '4. Hostinger Deployment Guide')
body('Your website is a React SPA built with Vite. The production build outputs to the /dist folder. Upload this to Hostinger.')

h(2, '4.1 Build the Project')
code_block([
    '# 1. Install dependencies',
    'npm install',
    '',
    '# 2. Create production build',
    'npm run build',
    '',
    '# Output is in the /dist folder — upload this to Hostinger',
])

h(2, '4.2 Hosting Option: Static Hosting (Recommended)')
bullet('Login to Hostinger hPanel at hpanel.hostinger.com')
bullet('Go to: Websites → Add Website → Static Website')
bullet('In File Manager: upload all contents of /dist to public_html folder')
bullet('Or use FTP: FileZilla → connect to your Hostinger FTP credentials')
bullet('CRITICAL: Upload the /dist CONTENTS (not the folder itself) into public_html')
bullet('Hostinger will auto-provision a free SSL certificate (Let\'s Encrypt) within minutes')

h(2, '4.3 Fix SPA Routing on Hostinger (Required)')
body('Without this file, refreshing /blog or /admin will show a 404 error. Create this file inside /dist before uploading:', bold=True)
body('Filename: .htaccess  (place inside public_html alongside index.html)')
code_block([
    'Options -MultiViews',
    'RewriteEngine On',
    'RewriteCond %{REQUEST_FILENAME} !-f',
    'RewriteRule ^ index.html [QSA,L]',
])
body('Alternatively in hPanel: Error Pages → 404 Error → set custom page to /index.html')

h(2, '4.4 Domain Setup')
steps = [
    '1. Buy domain on Hostinger: recommended igoproteincuts.in (Rs. 599/yr) or igoproteincuts.com',
    '2. In hPanel → Domains → the domain auto-connects if bought on Hostinger',
    '3. If domain is from another registrar: change nameservers to ns1.dns-parking.com & ns2.dns-parking.com',
    '4. In hPanel → SSL → Install Free SSL (Let\'s Encrypt) — takes 5-10 minutes',
    '5. Enable "Force HTTPS" in hPanel → SSL settings to redirect all HTTP traffic to HTTPS',
    '6. DNS propagation globally: 24–48 hours',
]
for s in steps:
    bullet(s)

h(2, '4.5 Environment Variables')
body('Your app uses Supabase + Gemini API keys. For static hosting, set these before building:', bold=True)
code_block([
    '# Create .env file in project root (NEVER commit to Git)',
    'VITE_SUPABASE_URL=https://your-project.supabase.co',
    'VITE_SUPABASE_ANON_KEY=your-anon-key-here',
    'VITE_GEMINI_API_KEY=your-gemini-api-key',
    '',
    '# Then run:  npm run build',
    '# The keys are baked into the /dist bundle at build time',
])

info_box('Pro Tip: Keep Vercel as Staging', [
    'Keep Vercel for auto-preview on every Git push (staging environment)',
    'Use Hostinger with your custom domain as the live production site',
    'Set up: hPanel → Domains → igoproteincuts.in points to Hostinger',
    '         Vercel → preview.igoproteincuts.in (subdomain CNAME to vercel)',
], bg='E8F5E9', border='2D7A3A')

doc.add_page_break()

# ════════════ SECTION 5 ════════════
h(1, '5. Specific Code Changes Required')

h(2, '5.1 Fix BrandNarrative Missing — src/App.tsx')
body('Find this block in the Home component and add one line:', bold=True)
code_block([
    '// BEFORE (current — BrandNarrative is missing):',
    '<Blog />',
    '<Newsletter />',
    '',
    '// AFTER (fix — add BrandNarrative before Newsletter):',
    '<Blog />',
    '<BrandNarrative />',
    '<Newsletter />',
])

h(2, '5.2 Fix Blank Scroll Area — src/index.css')
body('Add these CSS rules to fix the Lenis smooth scroll ghost height:', bold=True)
code_block([
    '/* Add to src/index.css */',
    'html.lenis, html.lenis body {',
    '  height: auto;',
    '}',
    '.lenis.lenis-smooth {',
    '  scroll-behavior: auto !important;',
    '}',
    '.lenis.lenis-smooth [data-lenis-prevent] {',
    '  overscroll-behavior: contain;',
    '}',
])

h(2, '5.3 Fix Admin Route Protection — src/App.tsx')
body('Wrap admin routes with an auth guard (create src/components/AdminGuard.tsx):', bold=True)
code_block([
    '// src/components/AdminGuard.tsx',
    'import { useEffect, useState } from "react";',
    'import { Navigate } from "react-router-dom";',
    'import { supabase } from "../lib/supabase";',
    '',
    'export default function AdminGuard({ children }) {',
    '  const [auth, setAuth] = useState(null);',
    '  useEffect(() => {',
    '    supabase.auth.getSession().then(({ data }) => {',
    '      setAuth(data.session?.user?.role === "admin" ? "ok" : "denied");',
    '    });',
    '  }, []);',
    '  if (!auth) return <LoadingFallback />;',
    '  if (auth === "denied") return <Navigate to="/admin/login" />;',
    '  return children;',
    '}',
])

h(2, '5.4 Add SEO Meta Tags — src/main.tsx + public/index.html')
code_block([
    '# Install react-helmet-async',
    'npm install react-helmet-async',
    '',
    '// In src/main.tsx — wrap app with HelmetProvider:',
    'import { HelmetProvider } from "react-helmet-async";',
    '<HelmetProvider><App /></HelmetProvider>',
    '',
    '// In sections/Hero.tsx — add Helmet:',
    'import { Helmet } from "react-helmet-async";',
    '<Helmet>',
    '  <title>IGO Protein Cuts — Farm-Fresh Meat Delivery in Coimbatore</title>',
    '  <meta name="description" content="Never frozen. Always fresh. Order chicken, mutton, fish delivered in 90 mins from heritage Tamil farms." />',
    '  <meta property="og:image" content="/og-image.jpg" />',
    '</Helmet>',
])

h(2, '5.5 Create public/robots.txt')
code_block([
    'User-agent: *',
    'Allow: /',
    'Disallow: /admin/',
    '',
    'Sitemap: https://igoproteincuts.in/sitemap.xml',
])

doc.add_paragraph()
doc.add_page_break()

# ════════════ SECTION 6 — ROADMAP ════════════
h(1, '6. Priority Roadmap Summary')

roadmap = [
    ('WEEK 1','IMMEDIATE FIXES','2D7A3A',[
        'Fix blank footer scroll bug (Lenis CSS)',
        'Add BrandNarrative to App.tsx Home',
        'Replace fake phone number in footer',
        'Fix dead Twitter / WhatsApp links',
        'Add custom 404 page',
        'Protect admin routes with auth guard',
        'Deploy to Hostinger (static build)',
    ]),
    ('WEEKS 2–4','CORE E-COMMERCE','1A6B4A',[
        'Integrate Razorpay payment gateway',
        'Build /checkout page (address + payment)',
        'Connect Supabase for orders storage',
        'Implement Supabase customer auth',
        'Move products from localStorage to Supabase',
        'Send order confirmation emails via Resend',
    ]),
    ('WEEKS 3–5','SEO & UX FIXES','2C5F8A',[
        'Add react-helmet-async meta tags',
        'Generate sitemap.xml + robots.txt',
        'Fix mobile layout on 375px screens',
        'Add WhatsApp Business chat widget',
        'Add Google Analytics 4',
        'Convert images to WebP, add lazy loading',
    ]),
    ('MONTH 2','ADVANCED FEATURES','5A3080',[
        'IGO Prime subscription billing (Razorpay)',
        'Real post-order reviews system',
        'B2B wholesale portal with login',
        'SMS order updates (MSG91 or Twilio)',
        'Multi-language (Tamil / English)',
        'PWA + offline service worker',
    ]),
]

for week, label, color_hex, tasks in roadmap:
    tbl_r = doc.add_table(rows=1, cols=2)
    tbl_r.alignment = WD_TABLE_ALIGNMENT.LEFT
    c0 = tbl_r.cell(0,0); c1 = tbl_r.cell(0,1)
    set_cell_bg(c0, color_hex); set_cell_bg(c1, 'F9FFF9')
    add_border(c0, color=color_hex); add_border(c1, color='CCCCCC')
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p0.add_run(week + '\n'); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = WHITE
    r2 = p0.add_run(label); r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0xA8,0xFF,0xC4)
    p0.paragraph_format.space_before = Pt(8); p0.paragraph_format.space_after = Pt(8)
    p1 = c1.paragraphs[0]
    for task in tasks:
        r = p1.add_run(f'  checkmark  {task}\n'); r.font.size = Pt(10)
    p1.paragraph_format.left_indent = Inches(0.05)
    p1.paragraph_format.space_before = Pt(5); p1.paragraph_format.space_after = Pt(5)
    doc.add_paragraph()

doc.add_page_break()

# ════════════ SECTION 7 ════════════
h(1, '7. Recommended Hostinger Setup & Costs')

plans = [
    ('Hosting Plan', 'Hostinger Business Web Hosting', 'Rs. 249/month (billed annually = Rs. 2,988/yr)'),
    ('Domain', 'igoproteincuts.in (recommended)', 'Rs. 599/year — buy on Hostinger for auto-connect'),
    ('SSL Certificate', 'Free Let\'s Encrypt via Hostinger', 'Included — auto-renewed every 90 days'),
    ('Business Email', 'support@igoproteincuts.in', 'Included with Business plan (up to 100 accounts)'),
    ('Database', 'Supabase (free tier)', 'Free up to 500MB + 50,000 monthly active users'),
    ('Payments', 'Razorpay', 'No monthly fee — 2% per successful transaction'),
    ('CDN', 'Cloudflare (free plan)', 'Add after setup — speeds up site globally + free DDoS protection'),
    ('Email Delivery', 'Resend (already in package.json)', 'Free up to 3,000 emails/month'),
    ('Analytics', 'Google Analytics 4', 'Free — add GA4 tag to index.html'),
]

tbl_p = doc.add_table(rows=len(plans)+1, cols=3)
tbl_p.alignment = WD_TABLE_ALIGNMENT.LEFT
for ci, ht in enumerate(['Component','Recommended Choice','Cost / Notes']):
    cell = tbl_p.cell(0, ci)
    set_cell_bg(cell, '2D7A3A'); add_border(cell, color='2D7A3A')
    p = cell.paragraphs[0]; r = p.add_run(ht)
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    p.paragraph_format.left_indent = Inches(0.05)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
for ri, (comp, choice, cost) in enumerate(plans, 1):
    for ci, text in enumerate([comp, choice, cost]):
        cell = tbl_p.cell(ri, ci)
        set_cell_bg(cell, 'FFFFFF' if ri%2==0 else 'F5F5F5'); add_border(cell, color='DDDDDD', sz=4)
        p = cell.paragraphs[0]; r = p.add_run(text)
        r.font.size = Pt(9.5); r.font.bold = (ci==0)
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# ── Footer ──
divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('IGO Protein Cuts — Website Audit Report  |  Confidential  |  May 2026  |  Prepared for: igobackend3@gmail.com')
r.font.size = Pt(8); r.font.color.rgb = GREY

# ── Save ──
out = "/sessions/confident-amazing-maxwell/mnt/Igo-Protein Cuts/IGO_Website_Audit_Report.docx"
doc.save(out)
print("Saved:", out)
